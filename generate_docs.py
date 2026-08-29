import sys
import os

try:
    import docx
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("docx module not found. Please install python-docx.")
    sys.exit(1)

def create_docx():
    doc = docx.Document()
    
    # Define styles and fonts
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("EXPERIMENT - 1\n")
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0, 51, 102)
    
    subtitle = title.add_run("AVL TREE CONSTRUCTION, INSERTION, AND DELETION")
    subtitle.bold = True
    subtitle.font.size = Pt(14)
    subtitle.font.color.rgb = RGBColor(102, 102, 102)
    
    doc.add_paragraph("\n")
    
    # Section: Aim
    p_aim_lbl = doc.add_paragraph()
    run = p_aim_lbl.add_run("AIM:")
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0, 51, 102)
    
    p_aim = doc.add_paragraph(
        "To write a C program to construct an AVL tree for a given set of elements, "
        "and implement insert and delete operations on the constructed tree. "
        "Write the contents of the tree into a new file using in-order traversal."
    )
    p_aim.paragraph_format.left_indent = Inches(0.25)
    
    # Section: Description
    p_desc_lbl = doc.add_paragraph()
    run = p_desc_lbl.add_run("DESCRIPTION:")
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0, 51, 102)
    
    p_desc = doc.add_paragraph(
        "An AVL tree (named after inventors Adelson-Velsky and Landis) is a self-balancing binary search tree (BST). "
        "In an AVL tree, the heights of the two child subtrees of any node differ by at most one. "
        "If at any time they differ by more than one, rebalancing is done to restore this property.\n\n"
        "Key concepts of AVL Trees:\n"
        "1. Balance Factor (BF): For any node N, BF(N) = Height(Left_Subtree(N)) - Height(Right_Subtree(N)). In a valid AVL tree, BF(N) must be in {-1, 0, 1}.\n"
        "2. Rotations: When insertion or deletion causes the balance factor of a node to deviate from this range, rotations are performed to rebalance the tree:\n"
        "   - Left-Left (LL) Case: Single right rotation around the unbalanced node.\n"
        "   - Right-Right (RR) Case: Single left rotation around the unbalanced node.\n"
        "   - Left-Right (LR) Case: Left rotation on the left child, followed by a right rotation on the unbalanced node.\n"
        "   - Right-Left (RL) Case: Right rotation on the right child, followed by a left rotation on the unbalanced node."
    )
    p_desc.paragraph_format.left_indent = Inches(0.25)
    
    # Section: Algorithm
    p_alg_lbl = doc.add_paragraph()
    run = p_alg_lbl.add_run("ALGORITHM:")
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0, 51, 102)
    
    p_alg = doc.add_paragraph(
        "1. Start the program.\n"
        "2. Define the structure for the tree node including standard left, right child pointers, data, and height variable.\n"
        "3. Implement utility functions: height() to get a node's height, getBalance() to compute balance factor, createNode() to allocate memory for a new node, and max() to find the maximum of two heights.\n"
        "4. Implement rotations:\n"
        "   - leftRotate(x): Rotates a node x to the left, updating parent/child links and heights.\n"
        "   - rightRotate(y): Rotates a node y to the right, updating links and heights.\n"
        "5. Implement insert(root, key):\n"
        "   - Recursively insert key as in standard BST.\n"
        "   - Update node height.\n"
        "   - Compute balance factor.\n"
        "   - Balance the node using LL, RR, LR, or RL rotations depending on balance factor and keys.\n"
        "6. Implement deleteNode(root, key):\n"
        "   - Perform standard BST deletion.\n"
        "   - Update ancestor heights.\n"
        "   - Check balance factors and rotate if unbalanced.\n"
        "7. Implement inOrderTraversal(root) to print node contents recursively (left, root, right).\n"
        "8. In the main() function, provide a menu-driven interface to insert, delete, traverse, or exit.\n"
        "9. End."
    )
    p_alg.paragraph_format.left_indent = Inches(0.25)
    
    # Section: Source Code
    p_src_lbl = doc.add_paragraph()
    run = p_src_lbl.add_run("SOURCE CODE:")
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0, 51, 102)
    
    # Read C code from avl_tree.c
    with open('avl_tree.c', 'r') as f:
        c_code = f.read()
        
    p_code = doc.add_paragraph()
    p_code.paragraph_format.left_indent = Inches(0.25)
    run = p_code.add_run(c_code)
    run.font.name = 'Consolas'
    run.font.size = Pt(9.5)
    
    # Section: Sample Output
    p_out_lbl = doc.add_paragraph()
    run = p_out_lbl.add_run("SAMPLE OUTPUT:")
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0, 51, 102)
    
    sample_output = (
        "AVL Tree Operations:\n"
        "1. Insert a node\n"
        "2. Delete a node\n"
        "3. In-order Traversal\n"
        "4. Exit\n"
        "Enter your choice: 1\n"
        "Enter the key to insert: 55\n\n"
        "AVL Tree Operations:\n"
        "1. Insert a node\n"
        "2. Delete a node\n"
        "3. In-order Traversal\n"
        "4. Exit\n"
        "Enter your choice: 3\n"
        "In-order Traversal: 55\n\n"
        "AVL Tree Operations:\n"
        "1. Insert a node\n"
        "2. Delete a node\n"
        "3. In-order Traversal\n"
        "4. Exit\n"
        "Enter your choice: 1\n"
        "Enter the key to insert: 85\n\n"
        "AVL Tree Operations:\n"
        "1. Insert a node\n"
        "2. Delete a node\n"
        "3. In-order Traversal\n"
        "4. Exit\n"
        "Enter your choice: 1\n"
        "Enter the key to insert: 74\n\n"
        "AVL Tree Operations:\n"
        "1. Insert a node\n"
        "2. Delete a node\n"
        "3. In-order Traversal\n"
        "4. Exit\n"
        "Enter your choice: 3\n"
        "In-order Traversal: 55 74 85\n\n"
        "AVL Tree Operations:\n"
        "1. Insert a node\n"
        "2. Delete a node\n"
        "3. In-order Traversal\n"
        "4. Exit\n"
        "Enter your choice: 2\n"
        "Enter the key to delete: 74\n\n"
        "AVL Tree Operations:\n"
        "1. Insert a node\n"
        "2. Delete a node\n"
        "3. In-order Traversal\n"
        "4. Exit\n"
        "Enter your choice: 3\n"
        "In-order Traversal: 55 85\n\n"
        "AVL Tree Operations:\n"
        "1. Insert a node\n"
        "2. Delete a node\n"
        "3. In-order Traversal\n"
        "4. Exit\n"
        "Enter your choice: 4\n"
        "Exiting..."
    )
    
    p_out = doc.add_paragraph()
    p_out.paragraph_format.left_indent = Inches(0.25)
    run = p_out.add_run(sample_output)
    run.font.name = 'Consolas'
    run.font.size = Pt(9.5)
    
    doc.save("adsa1.docx")
    print("adsa1.docx generated successfully!")

def create_html():
    with open('avl_tree.c', 'r') as f:
        c_code = f.read()
        
    html_content = f"""<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Experiment 1: AVL Tree</title>
    <style>
        body {{
            font-family: 'Inter', -apple-system, BlinkMacSystemFont, "Segoe UI", Roboto, Helvetica, Arial, sans-serif;
            line-height: 1.6;
            color: #333;
            max-width: 900px;
            margin: 0 auto;
            padding: 40px 20px;
            background-color: #f8fafc;
        }}
        .card {{
            background: #ffffff;
            padding: 40px;
            border-radius: 12px;
            box-shadow: 0 4px 6px -1px rgba(0, 0, 0, 0.1), 0 2px 4px -1px rgba(0, 0, 0, 0.06);
            border: 1px solid #e2e8f0;
        }}
        h1 {{
            color: #0f172a;
            font-size: 2.25rem;
            margin-bottom: 0.5rem;
            text-align: center;
        }}
        .subtitle {{
            text-align: center;
            color: #64748b;
            font-size: 1.1rem;
            margin-bottom: 2rem;
            text-transform: uppercase;
            letter-spacing: 1px;
            font-weight: 600;
        }}
        h2 {{
            color: #1e3a8a;
            border-bottom: 2px solid #e2e8f0;
            padding-bottom: 8px;
            margin-top: 2rem;
            font-size: 1.5rem;
        }}
        pre {{
            background-color: #0f172a;
            color: #f8fafc;
            padding: 20px;
            border-radius: 8px;
            overflow-x: auto;
            font-family: 'Fira Code', 'Consolas', monospace;
            font-size: 0.9rem;
        }}
        code {{
            font-family: 'Fira Code', 'Consolas', monospace;
        }}
        ol {{
            padding-left: 20px;
        }}
        li {{
            margin-bottom: 8px;
        }}
        .download-btn {{
            display: inline-block;
            background-color: #2563eb;
            color: white;
            padding: 12px 24px;
            text-decoration: none;
            border-radius: 6px;
            font-weight: bold;
            margin-top: 20px;
            transition: background-color 0.2s;
        }}
        .download-btn:hover {{
            background-color: #1d4ed8;
        }}
        .header-container {{
            text-align: center;
            margin-bottom: 30px;
        }}
    </style>
</head>
<body>
    <div class="card">
        <div class="header-container">
            <h1>EXPERIMENT - 1</h1>
            <div class="subtitle">AVL Tree Construction, Insertion, and Deletion</div>
        </div>

        <h2>AIM</h2>
        <p>To write a C program to construct an AVL tree for a given set of elements, and implement insert and delete operations on the constructed tree. Write the contents of the tree into a new file using in-order traversal.</p>

        <h2>DESCRIPTION</h2>
        <p>An AVL tree (named after inventors Adelson-Velsky and Landis) is a self-balancing binary search tree (BST). In an AVL tree, the heights of the two child subtrees of any node differ by at most one. If at any time they differ by more than one, rebalancing is done to restore this property.</p>
        <p>Key concepts of AVL Trees:</p>
        <ul>
            <li><strong>Balance Factor (BF):</strong> For any node N, <code>BF(N) = Height(Left_Subtree(N)) - Height(Right_Subtree(N))</code>. In a valid AVL tree, BF(N) must be in <code>{{-1, 0, 1}}</code>.</li>
            <li><strong>Rotations:</strong> When insertion or deletion causes the balance factor of a node to deviate from this range, rotations are performed to rebalance the tree:
                <ul>
                    <li><strong>Left-Left (LL) Case:</strong> Single right rotation around the unbalanced node.</li>
                    <li><strong>Right-Right (RR) Case:</strong> Single left rotation around the unbalanced node.</li>
                    <li><strong>Left-Right (LR) Case:</strong> Left rotation on the left child, followed by a right rotation on the unbalanced node.</li>
                    <li><strong>Right-Left (RL) Case:</strong> Right rotation on the right child, followed by a left rotation on the unbalanced node.</li>
                </ul>
            </li>
        </ul>

        <h2>ALGORITHM</h2>
        <ol>
            <li>Start the program.</li>
            <li>Define the structure for the tree node including standard left, right child pointers, data, and height variable.</li>
            <li>Implement utility functions: <code>height()</code> to get a node's height, <code>getBalance()</code> to compute balance factor, <code>createNode()</code> to allocate memory for a new node, and <code>max()</code> to find the maximum of two heights.</li>
            <li>Implement rotations:
                <ul>
                    <li><code>leftRotate(x)</code>: Rotates a node x to the left, updating parent/child links and heights.</li>
                    <li><code>rightRotate(y)</code>: Rotates a node y to the right, updating links and heights.</li>
                </ul>
            </li>
            <li>Implement <code>insert(root, key)</code>:
                <ul>
                    <li>Recursively insert key as in standard BST.</li>
                    <li>Update node height.</li>
                    <li>Compute balance factor.</li>
                    <li>Balance the node using LL, RR, LR, or RL rotations depending on balance factor and keys.</li>
                </ul>
            </li>
            <li>Implement <code>deleteNode(root, key)</code>:
                <ul>
                    <li>Perform standard BST deletion.</li>
                    <li>Update ancestor heights.</li>
                    <li>Check balance factors and rotate if unbalanced.</li>
                </ul>
            </li>
            <li>Implement <code>inOrderTraversal(root)</code> to print node contents recursively (left, root, right).</li>
            <li>In the <code>main()</code> function, provide a menu-driven interface to insert, delete, traverse, or exit.</li>
            <li>End.</li>
        </ol>

        <h2>SOURCE CODE</h2>
        <pre><code>{c_code.replace('<', '&lt;').replace('>', '&gt;')}</code></pre>

        <h2>SAMPLE OUTPUT</h2>
        <pre><code>AVL Tree Operations:
1. Insert a node
2. Delete a node
3. In-order Traversal
4. Exit
Enter your choice: 1
Enter the key to insert: 55

AVL Tree Operations:
1. Insert a node
2. Delete a node
3. In-order Traversal
4. Exit
Enter your choice: 3
In-order Traversal: 55

AVL Tree Operations:
1. Insert a node
2. Delete a node
3. In-order Traversal
4. Exit
Enter your choice: 1
Enter the key to insert: 85

AVL Tree Operations:
1. Insert a node
2. Delete a node
3. In-order Traversal
4. Exit
Enter your choice: 1
Enter the key to insert: 74

AVL Tree Operations:
1. Insert a node
2. Delete a node
3. In-order Traversal
4. Exit
Enter your choice: 3
In-order Traversal: 55 74 85

AVL Tree Operations:
1. Insert a node
2. Delete a node
3. In-order Traversal
4. Exit
Enter your choice: 2
Enter the key to delete: 74

AVL Tree Operations:
1. Insert a node
2. Delete a node
3. In-order Traversal
4. Exit
Enter your choice: 3
In-order Traversal: 55 85

AVL Tree Operations:
1. Insert a node
2. Delete a node
3. In-order Traversal
4. Exit
Enter your choice: 4
Exiting...</code></pre>

        <div style="text-align: center;">
            <a href="adsa1.docx" class="download-btn">Download Word Document (docx)</a>
        </div>
    </div>
</body>
</html>
"""
    with open("adsa1.html", "w") as f:
        f.write(html_content)
    print("adsa1.html generated successfully!")

if __name__ == "__main__":
    create_docx()
    create_html()
